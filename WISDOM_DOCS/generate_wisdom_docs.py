#!/usr/bin/env python3
"""Generate WISDOM Blueprint.docx + Prompts.docx (separate files)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

BASE = '/Users/khathahat/1.vibe code/10.Project-PROMETHEUS/WISDOM_DOCS'

def make_style(doc):
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(10.5)

def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x00,0xF0,0xFF) if level==1 else RGBColor(0xB0,0x26,0xFF) if level==2 else RGBColor(0xFF,0xAA,0x00)
    return h

def add_p(doc, text, bold=False, italic=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    if align: p.alignment = align
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_md_file(doc, filepath, title):
    """Read a .md file and add it to the document."""
    add_h(doc, title, 1)
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Headers
        if stripped.startswith('# ') and not stripped.startswith('# '):
            continue  # skip top-level (already added as title)
        if stripped.startswith('## '):
            add_h(doc, stripped[3:], 2)
        elif stripped.startswith('### '):
            add_h(doc, stripped[4:], 3)
        elif stripped.startswith('#### '):
            add_p(doc, stripped[5:], bold=True, size=11)
        # Code blocks
        elif stripped.startswith('```'):
            continue
        # Table rows
        elif stripped.startswith('|'):
            # Simple: add as code-like text
            add_code(doc, stripped)
        # List items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.':
            doc.add_paragraph(stripped[3:], style='List Number')
        # YAML/code-like content
        elif stripped.startswith('  ') and (':' in stripped or stripped.startswith('  -')):
            add_code(doc, stripped)
        # Regular text
        else:
            # Handle bold markers
            text = stripped.replace('**', '').replace('*', '')
            add_p(doc, text)

# ================================================================
# DOCUMENT 1: BLUEPRINT (all context files combined)
# ================================================================
print("Generating Blueprint.docx...")
doc = Document()
make_style(doc)

# Cover
doc.add_paragraph()
add_p(doc, 'PROJECT PROMETHEUS', bold=True, size=14, color=(0x00,0xF0,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'WISDOM', bold=True, size=36, color=(0xB0,0x26,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'AI-Generated Documentation LV.500', bold=True, size=14, color=(0xFF,0xAA,0x00), align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'Complete Blueprint — Context Package', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_p(doc, f'Generated: {datetime.date.today().strftime("%B %d, %Y")}', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'Origin: Project PROMETHEUS Invention #637', size=10, color=(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, '"Democratizing AI for 6.8 Billion People"', italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# TOC
add_h(doc, 'Table of Contents', 1)
toc = [
    'Part 1: PRD.md — Product Requirements Document',
    'Part 2: architecture.md — System Architecture',
    'Part 3: feature.json — Feature Specifications (25 features)',
    'Part 4: SKILL.md — Skill Tree & Capability Map (29 skills)',
]
for t in toc:
    add_p(doc, t, size=11)
doc.add_page_break()

# Part 1: PRD
add_md_file(doc, os.path.join(BASE, 'PRD.md'), 'Part 1: PRD — Product Requirements Document')
doc.add_page_break()

# Part 2: Architecture
add_md_file(doc, os.path.join(BASE, 'architecture.md'), 'Part 2: Architecture — System Design')
doc.add_page_break()

# Part 3: feature.json (formatted)
add_h(doc, 'Part 3: Feature Specifications', 1)
add_p(doc, '25 features across 4 phases. Full specifications in feature.json.', italic=True)
doc.add_paragraph()

import json
with open(os.path.join(BASE, 'feature.json'), 'r') as f:
    features = json.load(f)

for phase in features['phases']:
    add_h(doc, f"Phase {phase['phase']}: {phase['name']} ({phase['timeline']})", 2)
    add_p(doc, f"Codename: {phase['codename']}", italic=True)
    doc.add_paragraph()

    for feat in phase['features']:
        add_p(doc, f"[{feat['id']}] {feat['name']} — {feat['priority']}", bold=True, size=11)
        add_p(doc, f"Module: {feat['module']} | File: {feat['file']} | Effort: {feat['effort']}")
        add_p(doc, feat['description'])
        add_p(doc, 'Acceptance Criteria:', bold=True)
        for ac in feat['acceptance_criteria']:
            doc.add_paragraph(ac, style='List Bullet')
        add_p(doc, f"User Story: {feat['user_story']}", italic=True, size=10)
        doc.add_paragraph()

doc.add_page_break()

# Part 4: SKILL
add_md_file(doc, os.path.join(BASE, 'SKILL.md'), 'Part 4: SKILL — Capability Map')

# Save
blueprint_path = os.path.join(BASE, 'WISDOM_Blueprint_LV500.docx')
doc.save(blueprint_path)
print(f"Blueprint saved: {blueprint_path}")

# ================================================================
# DOCUMENT 2: PROMPTS (separate, copy-paste ready)
# ================================================================
print("\nGenerating Prompts.docx...")
doc2 = Document()
make_style(doc2)

# Cover
doc2.add_paragraph()
add_p(doc2, 'PROJECT PROMETHEUS WISDOM', bold=True, size=14, color=(0x00,0xF0,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc2, 'Claude Code Prompts', bold=True, size=28, color=(0xB0,0x26,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc2, 'Copy & Paste Ready — 10 Prompts', bold=True, size=14, color=(0xFF,0xAA,0x00), align=WD_ALIGN_PARAGRAPH.CENTER)
doc2.add_paragraph()
add_p(doc2, f'Generated: {datetime.date.today().strftime("%B %d, %Y")}', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc2, 'Use these prompts in order with Claude Code to build WISDOM from scratch.', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc2.add_page_break()

# Instructions
add_h(doc2, 'How to Use These Prompts', 1)
instructions = [
    "1. Create a new project folder: mkdir prometheus-wisdom && cd prometheus-wisdom",
    "2. Open Claude Code: claude",
    "3. Copy-paste Prompt 0 to initialize the project",
    "4. After each prompt completes, copy-paste the next prompt",
    "5. Prompts are ordered by dependency — follow the sequence",
    "6. Each prompt is self-contained — Claude Code has full context",
    "7. After Prompt 9, you'll have a fully working WISDOM!",
]
for inst in instructions:
    add_p(doc2, inst, size=11)

add_p(doc2, '\nPrerequisites:', bold=True, size=12)
prereqs = [
    "Python 3.12+ installed",
    "Ollama installed (brew install ollama && ollama pull llama3)",
    "OR Google AI API key (free: aistudio.google.com)",
    "Git installed",
    "Claude Code CLI installed",
]
for p in prereqs:
    doc2.add_paragraph(p, style='List Bullet')

doc2.add_page_break()

# All prompts
prompts = [
    {
        'num': 0,
        'title': 'Project Initialization & Structure',
        'desc': 'Creates the entire project skeleton, dependencies, and configuration.',
        'prompt': """Create a new Python project called "prometheus-wisdom" — an open-source AI companion for humanity.

CONTEXT: This is WISDOM from Project PROMETHEUS — an AI personal assistant designed to help the 84% of humanity (6.8 billion people) who have never used AI. It adapts to any language, any skill level, and grows with the user from AI beginner to software builder.

Create this exact project structure:

prometheus-wisdom/
├── wisdom/
│   ├── __init__.py              # Main WISDOM class with chat() method
│   ├── brain/
│   │   ├── __init__.py
│   │   ├── knowledge_graph.py   # Neo4j Aura (free) + SQLite fallback
│   │   ├── memory_manager.py    # Short-term (20 msgs) + Long-term (ChromaDB vectors)
│   │   ├── user_profile.py      # SQLite user profiles (name, lang, level, goals)
│   │   └── embeddings.py        # Ollama nomic-embed-text + Gemini fallback
│   ├── voice/
│   │   ├── __init__.py
│   │   ├── chat_engine.py       # LangChain conversation with streaming
│   │   ├── language_detect.py   # Auto-detect from Unicode + LLM confirmation
│   │   ├── prompt_templates.py  # WISDOM personality + multilingual templates
│   │   └── tone_adapter.py      # Auto-adjust complexity/formality per user
│   ├── heart/
│   │   ├── __init__.py
│   │   ├── privacy_manager.py   # PII sanitization (email, phone, CC → masked)
│   │   ├── federated_core.py    # Opt-in anonymous aggregate learning
│   │   ├── community_knowledge.py # Shared Q&A knowledge base
│   │   └── feedback_loop.py     # User ratings + improvement pipeline
│   ├── soul/
│   │   ├── __init__.py
│   │   ├── adaptation_engine.py # Core personalization logic
│   │   ├── skill_assessor.py    # 5-question adaptive assessment (0-10 scale)
│   │   ├── learning_path.py     # 7-level curriculum (AI basics → build AI tools)
│   │   └── goal_tracker.py      # Goals, milestones, achievements, badges
│   ├── body/
│   │   ├── __init__.py
│   │   ├── app.py               # Streamlit main app (chat UI + sidebar)
│   │   ├── api.py               # FastAPI REST API (/chat, /profile, /progress)
│   │   └── components/
│   │       ├── __init__.py
│   │       ├── chat.py          # Chat bubbles, streaming, onboarding flow
│   │       └── dashboard.py     # Progress bars, skill radar, badges
│   ├── core/
│   │   ├── __init__.py
│   │   ├── llm_provider.py      # Dual-path: Ollama (primary) + Gemini (fallback)
│   │   ├── config.py            # Env-based config with sensible defaults
│   │   ├── constants.py         # App constants, version, supported languages
│   │   └── orchestrator.py      # Conversation flow controller (all modules wired)
│   └── tests/
│       ├── __init__.py
│       ├── test_brain.py
│       ├── test_voice.py
│       ├── test_heart.py
│       ├── test_soul.py
│       └── test_body.py
├── .env.example                 # All env vars with comments
├── .gitignore                   # Python, venv, .env, *.db, __pycache__
├── .streamlit/
│   └── config.toml              # Dark theme, wide layout
├── requirements.txt             # All dependencies pinned
├── runtime.txt                  # python-3.12
├── Makefile                     # install, test, run, api, lint
├── setup.py                     # Package setup
└── LICENSE                      # MIT

requirements.txt should include:
langchain, langchain-community, langchain-google-genai, langchain-ollama,
chromadb, neo4j, streamlit, fastapi, uvicorn, python-dotenv, pytest,
httpx, plotly

.env.example should document:
GOOGLE_API_KEY, NEO4J_URI, NEO4J_PASSWORD, OLLAMA_BASE_URL, WISDOM_DB_PATH

Makefile targets: install, test, run, api, lint, clean

Initialize git repo with initial commit.
Make every __init__.py importable with proper __all__.
Add MIT LICENSE file."""
    },
    {
        'num': 1,
        'title': 'Core — LLM Provider (Dual-Path)',
        'desc': 'The foundation: Ollama (local free) + Gemini (cloud free) with auto-switching.',
        'prompt': """Build wisdom/core/llm_provider.py — the LLM abstraction layer.

REQUIREMENTS:
1. Class LLMProvider with methods:
   - get_llm() → LangChain BaseChatModel (ChatOllama or ChatGoogleGenerativeAI)
   - get_embeddings() → LangChain Embeddings (OllamaEmbeddings or GoogleGenerativeAIEmbeddings)
   - is_local() → bool
   - health_check() → dict with provider, model, latency_ms, status

2. Auto-detection logic:
   - Try Ollama at OLLAMA_BASE_URL (default localhost:11434) with 2s timeout
   - If Ollama responds: use ChatOllama(model="llama3") + OllamaEmbeddings(model="nomic-embed-text")
   - If Ollama fails: use ChatGoogleGenerativeAI(model="gemini-2.0-flash") + GoogleGenerativeAIEmbeddings
   - If both fail: raise clear error with setup instructions

3. Configuration (from env with defaults):
   - OLLAMA_BASE_URL = "http://localhost:11434"
   - OLLAMA_MODEL = "llama3"
   - OLLAMA_EMBED_MODEL = "nomic-embed-text"
   - GOOGLE_API_KEY = None (optional, for Gemini)
   - GEMINI_MODEL = "gemini-2.0-flash"

4. Streaming support: both providers must support streaming
5. Singleton pattern: one instance shared across the app
6. Logging: log which provider was selected and why

Also build wisdom/core/config.py:
- Load .env file
- Provide typed config values with defaults
- No config file needed for default setup (zero-config)

And wisdom/core/constants.py:
- VERSION = "1.0.0"
- APP_NAME = "PROMETHEUS WISDOM"
- SUPPORTED_LANGUAGES = {code: name} for 10 languages
- SKILL_LEVELS = {0-10 descriptions}
- LEARNING_LEVELS = {1-7 descriptions}

Write tests in test_brain.py (mock Ollama/Gemini connections)."""
    },
    {
        'num': 2,
        'title': 'Brain — User Profile & Memory',
        'desc': 'WISDOM remembers you: SQLite profiles, short-term + long-term memory, embeddings.',
        'prompt': """Build the wisdom/brain/ module — WISDOM's memory system.

FILE 1: wisdom/brain/user_profile.py
- Class UserProfile with SQLite backend (wisdom.db)
- Auto-create tables on first use
- Methods:
  - create_user(name, language) → user_id (UUID)
  - get_user(user_id) → dict
  - update_user(user_id, **kwargs)
  - delete_user(user_id) — right to be forgotten
  - export_profile(user_id) → JSON string
  - import_profile(json_string) → user_id
- Schema: id, name, language, skill_level (float 0-10), interests (JSON array),
  goals (JSON array), preferences (JSON dict), created_at, last_seen
- Auto-update last_seen on every get_user()

FILE 2: wisdom/brain/memory_manager.py
- Class MemoryManager
- Short-term memory:
  - In-memory list of last 20 messages [{role, content, timestamp}]
  - add_message(role, content)
  - get_history() → list (for LLM context)
  - get_history_text() → formatted string
  - When exceeding 20: summarize oldest 5 via LLM, keep summary + 15 newest
- Long-term memory:
  - Store important facts as vector embeddings in ChromaDB
  - save_fact(user_id, fact_text, metadata)
  - recall(user_id, query, top_k=5) → list of relevant facts
  - save_conversation_summary(user_id, summary)
  - Uses wisdom/brain/embeddings.py for vectors
- Session management:
  - on_session_start(user_id) — load last summary
  - on_session_end(user_id) — summarize and persist

FILE 3: wisdom/brain/embeddings.py
- Class EmbeddingManager
- Uses LLMProvider.get_embeddings() (Ollama or Gemini)
- Methods:
  - embed_text(text) → list[float] (768-dim)
  - embed_batch(texts) → list[list[float]]
  - similarity(text1, text2) → float (cosine)
- ChromaDB collection management:
  - get_or_create_collection(name)
  - add_documents(collection, texts, metadatas, ids)
  - query(collection, query_text, top_k) → results

FILE 4: wisdom/brain/knowledge_graph.py
- Class KnowledgeGraph
- Primary: Neo4j Aura (free tier) if NEO4J_URI set
- Fallback: SQLite-based graph (nodes + edges tables)
- Methods:
  - add_node(node_type, properties) → node_id
  - add_relationship(from_id, to_id, rel_type, properties)
  - get_user_knowledge(user_id) → list of topics learned
  - get_related_topics(topic_id) → connected topics
  - get_learning_recommendations(user_id) → next topics to learn
  - For SQLite fallback: nodes(id, type, data_json), edges(from_id, to_id, type, data_json)
- Auto-detect: try Neo4j first, fall back to SQLite silently

Write tests for all CRUD operations (mock DB connections)."""
    },
    {
        'num': 3,
        'title': 'Voice — Chat Engine & Multilingual NLP',
        'desc': 'How WISDOM speaks: language detection, adaptive prompts, streaming chat.',
        'prompt': """Build the wisdom/voice/ module — WISDOM's communication layer.

FILE 1: wisdom/voice/language_detect.py
- Function detect_language(text: str) → dict with:
  - code: "th", name: "Thai", confidence: 0.98, script: "Thai"
- Detection method:
  1. Unicode range check (fast, first pass):
     - Thai: U+0E00-U+0E7F
     - Devanagari (Hindi): U+0900-U+097F
     - Arabic: U+0600-U+06FF
     - CJK (Chinese): U+4E00-U+9FFF
     - Japanese Hiragana/Katakana: U+3040-U+30FF
     - Korean Hangul: U+AC00-U+D7A3
  2. If Latin script: use trigram analysis (en vs es vs pt vs fr vs sw vs id)
  3. Return with confidence score
- No external API calls needed (fast, offline)

FILE 2: wisdom/voice/prompt_templates.py
- WISDOM_SYSTEM_PROMPT — Main personality template with placeholders:
  {user_name}, {user_language}, {skill_level}, {current_goal},
  {learning_progress}, {retrieved_context}, {current_mode}
- Mode-specific additions:
  TEACHER_PROMPT — Patient, step-by-step, lots of examples
  RESEARCHER_PROMPT — Thorough, cite-style, comprehensive
  QUIZ_MASTER_PROMPT — Ask questions, grade, encourage
  CODE_HELPER_PROMPT — Write code, explain line by line
  FREE_CHAT_PROMPT — Natural, warm, open-ended
- WELCOME_PROMPT — First-time user onboarding
- BEGINNER_ADAPTER — "Explain like I'm 5" additions
- EXPERT_ADAPTER — "Full technical depth" additions
- All prompts work in any language (LLM translates naturally)

FILE 3: wisdom/voice/tone_adapter.py
- Class ToneAdapter
- Methods:
  - analyze_user_message(message, history) → dict with:
    estimated_level (1-10), emotional_state, formality_preference
  - get_adapted_prompt(base_prompt, user_profile, analysis) → str
  - Signals to detect:
    confusion: short answers, "?", "I don't get it"
    frustration: aggressive punctuation, "doesn't work"
    excitement: "wow", "cool", exclamation marks
    boredom: "ok", "sure", minimal responses
    advanced: technical terms, code snippets

FILE 4: wisdom/voice/chat_engine.py
- Class ChatEngine
- Uses LLMProvider, MemoryManager, ToneAdapter, PromptTemplates
- Methods:
  - chat(user_id, message, stream=True) → generator of response chunks
  - chat_sync(user_id, message) → full response string
  - set_mode(mode: str) — "teacher", "researcher", "quiz", "code", "chat"
  - get_mode() → current mode
- Flow per message:
  1. Detect language (if first message)
  2. Load user profile
  3. Retrieve relevant memory (RAG)
  4. Analyze tone/complexity
  5. Build system prompt (template + context + adaptation)
  6. Call LLM with streaming
  7. Save message + response to memory
  8. Update user profile (last_seen, etc.)

Write tests with mocked LLM responses."""
    },
    {
        'num': 4,
        'title': 'Soul — Skill Assessment & Learning Path',
        'desc': 'Personalization: assess user level, generate curriculum, track progress.',
        'prompt': """Build the wisdom/soul/ module — WISDOM's personalization engine.

FILE 1: wisdom/soul/skill_assessor.py
- Class SkillAssessor
- Methods:
  - start_assessment(user_id) → first question
  - answer_question(user_id, answer) → next question or results
  - get_results(user_id) → SkillProfile with scores per category
- 5 categories with 2 questions each (10 total, adaptive):
  1. AI Awareness (weight 0.30):
     Q1: "Have you heard of AI assistants like ChatGPT?" → yes/no/what's that
     Q2: "Name something AI can do" → free text, LLM grades
  2. Prompt Skills (weight 0.25):
     Q1: "How would you ask AI to help write an email?" → free text
     Q2: "What makes a good question for AI?" → free text
  3. Digital Literacy (weight 0.20):
     Q1: "Rate your comfort with computers (1-5)" → number
     Q2: "What apps do you use daily?" → free text
  4. Coding Ability (weight 0.15):
     Q1: "Have you ever written code?" → yes/some/no
     Q2: "What does 'if-else' mean?" → free text
  5. Domain Knowledge (weight 0.10):
     Q1: "What's your field of work/study?" → free text
     Q2: "How could AI help in your field?" → free text
- Adaptive: if Q1 answer indicates 0, skip Q2 (score 0)
- Score each category 0-10, compute weighted composite
- Store results in user profile

FILE 2: wisdom/soul/learning_path.py
- Class LearningPath
- 7 levels, each with 3-4 modules, each module has:
  - title, description, objectives (list), lessons (list of dicts),
    exercises (list), quiz_questions (list), estimated_time
- Methods:
  - generate_path(user_id) → PersonalizedPath based on assessment + interests
  - get_current_module(user_id) → current module with progress
  - complete_lesson(user_id, module_id, lesson_id) → updated progress
  - take_quiz(user_id, module_id) → quiz questions
  - grade_quiz(user_id, module_id, answers) → score + pass/fail
  - get_progress(user_id) → overall progress dict
- Content is generated by LLM, personalized to user's:
  - Language, interests, skill level, goals
  - Example: farmer gets farming examples, teacher gets education examples
- Level definitions:
  L1: "Hello AI" — What is AI, expectations, first conversation
  L2: "Talk to AI" — Prompt engineering basics, context, specificity
  L3: "AI Daily Life" — Productivity, learning, creativity, problem-solving
  L4: "How AI Thinks" — LLMs, training, bias, limitations, ethics
  L5: "Code with AI" — Python basics, AI-assisted coding, debugging
  L6: "Build with AI" — Web apps, APIs, testing, deployment
  L7: "Create AI Tools" — Build AI-powered apps, open source

FILE 3: wisdom/soul/goal_tracker.py
- Class GoalTracker
- Methods:
  - set_goal(user_id, description) → goal_id + auto-generated milestones
  - get_goals(user_id) → list of goals with progress
  - update_milestone(user_id, goal_id, milestone_id) → updated progress
  - check_achievements(user_id) → list of newly earned badges
  - get_achievements(user_id) → all earned badges
- WISDOM auto-breaks goals into 3-5 milestones using LLM
- Achievement badges:
  FIRST_CONTACT: First conversation
  CURIOUS_MIND: Asked 50 questions
  LEVEL_UP: Completed any learning module
  PERFECT_SCORE: 100% on a quiz
  STREAK_7: 7-day usage streak
  CODE_ROOKIE: First code written
  BUILDER: Completed Level 6
  CREATOR: Completed Level 7
  HELPER: Contributed to community

FILE 4: wisdom/soul/adaptation_engine.py
- Class AdaptationEngine
- Methods:
  - adapt(user_id, message, history) → AdaptationResult with:
    recommended_mode, difficulty_level, prompt_modifiers, suggested_topic
  - Combines: SkillAssessor + GoalTracker + UserProfile + ConversationHistory
  - Logic:
    New user → Welcome mode
    Low skill + learning path active → Teacher mode
    User asks factual question → Researcher mode
    End of module → Quiz mode
    User shares code → Code helper mode
    Default → Free chat mode

Write tests for assessment flow, learning path generation, and goal tracking."""
    },
    {
        'num': 5,
        'title': 'Heart — Privacy & Community',
        'desc': 'Protecting users: PII sanitization, federated learning, community knowledge.',
        'prompt': """Build the wisdom/heart/ module — WISDOM's privacy and community layer.

FILE 1: wisdom/heart/privacy_manager.py
- Class PrivacyManager
- Methods:
  - sanitize(text: str) → str with PII masked
  - detect_pii(text: str) → list of PII findings with type and position
  - is_safe_for_cloud(text: str) → bool
- PII patterns to detect and mask:
  - Email addresses → [EMAIL]
  - Phone numbers (international formats) → [PHONE]
  - Credit card numbers → [CARD]
  - Thai national ID (13 digits) → [THAI_ID]
  - US SSN → [SSN]
  - Physical addresses (heuristic) → [ADDRESS]
  - IP addresses → [IP]
- When applied:
  - ALWAYS before sending to Gemini (cloud)
  - NEVER for Ollama (local) — data stays on device
  - ALWAYS before community knowledge sharing
- User can toggle on/off per session
- Logging: record what was sanitized (without the actual data)

FILE 2: wisdom/heart/feedback_loop.py
- Class FeedbackLoop
- Methods:
  - rate_response(user_id, message_id, rating: 1-5, comment: str)
  - get_ratings(user_id) → list of ratings
  - get_average_rating() → float
  - get_improvement_suggestions() → based on low-rated responses
- Store in SQLite feedback table
- Used to improve prompt templates over time

FILE 3: wisdom/heart/community_knowledge.py
- Class CommunityKnowledge
- Methods:
  - share_qa(question, answer, user_id) → qa_id (anonymous)
  - search_community(query, top_k=5) → relevant community Q&As
  - upvote(qa_id) / downvote(qa_id)
  - get_popular(category, limit=10) → top voted Q&As
  - report(qa_id, reason) — flag inappropriate content
- All submissions are anonymized (no user_id stored with content)
- Privacy manager applied before storage
- Stored in shared SQLite or remote DB (configurable)

FILE 4: wisdom/heart/federated_core.py
- Class FederatedCore
- Methods:
  - opt_in(user_id) / opt_out(user_id) → toggle participation
  - is_opted_in(user_id) → bool
  - collect_local_metrics(user_id) → anonymous aggregate data:
    * Most asked topic categories (not actual questions)
    * Average difficulty level
    * Common confusion points (topic, not content)
    * Module completion rates
  - apply_differential_privacy(metrics) → noised metrics
  - preview_shared_data(user_id) → show user what would be shared
- Differential privacy: add Laplacian noise (epsilon=1.0)
- NEVER shares: actual messages, names, PII, specific questions
- Phase 4 feature — stub methods for now, implement later

Write tests for PII detection (with various formats) and sanitization."""
    },
    {
        'num': 6,
        'title': 'Body — Streamlit Chat Interface',
        'desc': 'The face of WISDOM: beautiful chat UI, sidebar, onboarding, mobile-friendly.',
        'prompt': """Build wisdom/body/app.py — the main Streamlit web application.

REQUIREMENTS:
1. Page config:
   - Title: "WISDOM — AI Companion"
   - Icon: robot emoji
   - Layout: wide
   - Dark theme via .streamlit/config.toml

2. Session state management:
   - user_id: persistent across reloads (st.session_state)
   - chat_history: list of messages
   - current_mode: active WISDOM mode
   - If no user_id: show onboarding flow

3. Onboarding flow (first-time users):
   - WISDOM introduces itself in a welcoming way
   - Detect language from user's first response
   - Ask for name (optional)
   - Offer 3 paths: [Learn about AI] [Just Chat] [Help me with something]
   - Create user profile, store in session state
   - Transition to main chat

4. Main chat interface:
   - st.chat_message for WISDOM (avatar: robot) and user (avatar: person)
   - st.chat_input at bottom
   - Streaming responses with st.write_stream
   - Markdown rendering in responses
   - Code blocks with syntax highlighting (via markdown)
   - Auto-scroll to latest message

5. Sidebar:
   - User name + current level display
   - Learning progress bar (if on learning path)
   - Quick action buttons:
     * "Teach me" → teacher mode
     * "Quiz me" → quiz mode
     * "Help me code" → code mode
     * "My progress" → show dashboard
   - Settings expander:
     * Language selector dropdown
     * Difficulty slider (1-10)
   - Footer: "Open Source | Made for Humanity"

6. Dashboard (in sidebar or separate page):
   - Learning progress per module (progress bars)
   - Skill radar chart (plotly, 5 categories)
   - Achievement badges (emoji + name)
   - Total conversations count
   - Days active

7. Mobile-responsive:
   - Large fonts (16px min)
   - Touch-friendly buttons
   - Sidebar collapses on mobile
   - Input always visible at bottom

Also create wisdom/body/components/chat.py:
   - render_message(role, content) — styled chat bubble
   - render_onboarding() — welcome flow
   - render_streaming(generator) — streaming display

And wisdom/body/components/dashboard.py:
   - render_progress(user_id) — progress bars + radar chart
   - render_achievements(user_id) — badge display
   - render_stats(user_id) — conversation stats

The UI should be so intuitive that someone who has never used a computer
can understand it. Big text, clear buttons, no jargon."""
    },
    {
        'num': 7,
        'title': 'Body — FastAPI REST API',
        'desc': 'API endpoints for third-party integrations and mobile apps.',
        'prompt': """Build wisdom/body/api.py — FastAPI REST API for WISDOM.

ENDPOINTS:

POST /api/v1/chat
  Body: {"user_id": "abc", "message": "Hello", "stream": false}
  Response: {"response": "...", "language": "th", "mode": "chat"}
  If stream=true: return Server-Sent Events (SSE)

GET /api/v1/profile/{user_id}
  Response: {"name": "...", "language": "th", "skill_level": 3.5, ...}

PUT /api/v1/profile/{user_id}
  Body: {"language": "en", "preferences": {...}}
  Response: {"status": "updated"}

POST /api/v1/profile
  Body: {"name": "...", "language": "th"}
  Response: {"user_id": "...", "status": "created"}

DELETE /api/v1/profile/{user_id}
  Response: {"status": "deleted"} (right to be forgotten)

GET /api/v1/progress/{user_id}
  Response: {"level": 3, "modules": [...], "score": 65, "badges": [...]}

GET /api/v1/learning-path/{user_id}
  Response: {"current_level": 3, "next_module": {...}, "progress": 0.6}

POST /api/v1/feedback
  Body: {"user_id": "abc", "rating": 5, "comment": "Great!"}
  Response: {"status": "recorded"}

GET /api/v1/health
  Response: {"status": "ok", "llm_provider": "ollama", "version": "1.0.0", "uptime": 3600}

REQUIREMENTS:
- CORS middleware (allow all origins for now)
- Request validation with Pydantic models
- Error handling with proper HTTP status codes
- Rate limiting: 100 requests/minute per user_id
- OpenAPI/Swagger docs auto-generated at /docs
- Can run alongside Streamlit or independently
- Entry point: uvicorn wisdom.body.api:app --port 8000

Write tests using httpx AsyncClient for all endpoints."""
    },
    {
        'num': 8,
        'title': 'Core — Orchestrator (Wire Everything Together)',
        'desc': 'Connect all modules: Brain + Voice + Soul + Heart + Body into one flow.',
        'prompt': """Build wisdom/core/orchestrator.py — the conversation flow controller.

This is the CENTRAL NERVOUS SYSTEM that wires all modules together.

Class WISDOMOrchestrator:
  - __init__(): Initialize all modules:
    * LLMProvider (core)
    * UserProfile (brain)
    * MemoryManager (brain)
    * KnowledgeGraph (brain)
    * ChatEngine (voice)
    * LanguageDetect (voice)
    * ToneAdapter (voice)
    * AdaptationEngine (soul)
    * SkillAssessor (soul)
    * LearningPath (soul)
    * GoalTracker (soul)
    * PrivacyManager (heart)
    * FeedbackLoop (heart)

  - chat(user_id: str, message: str, stream: bool = True):
    FULL ORCHESTRATION FLOW:
    1. Get or create user profile
    2. Detect language (if not set)
    3. Sanitize message for PII (if cloud LLM)
    4. Retrieve relevant long-term memory (RAG, top 5)
    5. Get knowledge graph context (topics learned, goals)
    6. Run adaptation engine (determine mode, difficulty, prompt mods)
    7. Build system prompt (template + context + adaptations)
    8. Call LLM via ChatEngine (with streaming)
    9. Save user message + WISDOM response to short-term memory
    10. Extract and save key facts to long-term memory
    11. Update knowledge graph (new topics discussed)
    12. Update user profile (last_seen, interaction count)
    13. Return response (generator if streaming, string if not)

  - start_assessment(user_id) → first assessment question
  - answer_assessment(user_id, answer) → next question or results
  - get_progress(user_id) → full progress report
  - set_goal(user_id, goal) → created goal with milestones
  - export_user_data(user_id) → complete JSON export
  - delete_user_data(user_id) → complete data deletion

Also update wisdom/__init__.py:
  - from wisdom.core.orchestrator import WISDOMOrchestrator as WISDOM
  - So users can do: from wisdom import WISDOM; j = WISDOM(); j.chat(...)

Make sure all error handling is graceful:
  - If any module fails, others still work
  - If Neo4j down → SQLite graph
  - If Ollama down → Gemini
  - If memory fails → continue without context
  - Log all errors but never crash

Write integration tests that test the full flow with mocked LLM."""
    },
    {
        'num': 9,
        'title': 'Final — Testing, README, Deploy',
        'desc': 'Polish everything: 50+ tests, documentation, and deploy to Streamlit Cloud.',
        'prompt': """Final polish for PROMETHEUS WISDOM — testing, docs, and deployment.

1. TESTING (wisdom/tests/):
   Run all existing tests and fix any failures.
   Add missing tests to reach 50+ total:
   - test_brain.py: user CRUD, memory store/recall, embeddings mock, KG operations
   - test_voice.py: language detection (Thai, English, Hindi, Spanish, Arabic, Chinese),
     prompt template rendering, tone analysis
   - test_heart.py: PII detection (email, phone, CC, Thai ID), sanitization,
     feedback recording
   - test_soul.py: assessment flow, learning path generation, goal creation,
     badge earning
   - test_body.py: API endpoints (all 8), response formats, error handling
   All tests must pass: pytest wisdom/tests/ -v

2. README.md:
   - Project name, tagline, vision (1 paragraph)
   - The Problem: 84% of humanity has never used AI (reference the dot visualization)
   - The Solution: WISDOM — personal AI companion
   - Quick Start (3 commands):
     ```
     git clone ... && cd prometheus-wisdom
     make install
     make run
     ```
   - Screenshots placeholder
   - Architecture overview (5 modules: Brain, Voice, Heart, Soul, Body)
   - Tech stack table
   - API documentation link (/docs)
   - Contributing guide (fork, branch, PR)
   - License: MIT
   - Credits: Project PROMETHEUS

3. DEPLOYMENT:
   - Verify .streamlit/config.toml (dark theme, wide layout)
   - Verify runtime.txt (python-3.12)
   - Verify requirements.txt (all deps, no version conflicts)
   - Add Procfile for alternative hosting: web: streamlit run wisdom/body/app.py
   - Verify .gitignore covers: .env, *.db, venv/, __pycache__/, .chroma/

4. FINAL VERIFICATION:
   - Run: make install && make test
   - Run: make run (verify Streamlit starts)
   - Run: make api (verify FastAPI starts)
   - Verify chat works end-to-end
   - Verify onboarding flow works for new users

The goal: anyone in the world can clone this repo and have their own
WISDOM running in under 5 minutes."""
    },
]

for p_data in prompts:
    add_h(doc2, f"Prompt {p_data['num']}: {p_data['title']}", 1)
    add_p(doc2, p_data['desc'], italic=True, size=11)
    add_p(doc2, 'Copy everything below this line into Claude Code:', bold=True, size=10, color=(0xFF,0xAA,0x00))
    doc2.add_paragraph()

    # Add prompt as code block, split into readable chunks
    lines = p_data['prompt'].strip().split('\n')
    for line in lines:
        add_code(doc2, line)

    doc2.add_page_break()

# Final page
add_h(doc2, 'Quick Reference — Prompt Order', 1)
for p in prompts:
    add_p(doc2, f"Prompt {p['num']}: {p['title']}", bold=True)
    add_p(doc2, p['desc'], size=10)
    doc2.add_paragraph()

add_p(doc2, '\nTotal: 10 prompts → Complete WISDOM from scratch', bold=True, size=12, color=(0x00,0xF0,0xFF))
add_p(doc2, 'Estimated time: 4-6 hours with Claude Code', size=11)
add_p(doc2, '\n"Every grey dot deserves a WISDOM."', italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

prompts_path = os.path.join(BASE, 'WISDOM_Prompts_LV500.docx')
doc2.save(prompts_path)
print(f"Prompts saved: {prompts_path}")

print("\n✅ All documents generated!")
print(f"  Blueprint: {blueprint_path}")
print(f"  Prompts:   {prompts_path}")
